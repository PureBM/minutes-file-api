from flask import Flask, request, jsonify
from datetime import date
import csv, io, base64
from docx import Document

app = Flask(__name__)

@app.post("/api/create_minutes")
def create_minutes():
    data = request.get_json()
    minutes_text = data.get("minutes", "")
    actions = data.get("actions", [])
    title = data.get("title", "Meeting")  # default if not provided


    # Build Word file
    doc = Document()
    doc.add_heading(f"{title} – {date.today()}", 0)
    doc.add_paragraph(minutes_text)
    word_stream = io.BytesIO()
    doc.save(word_stream)
    word_b64 = base64.b64encode(word_stream.getvalue()).decode("utf-8")

    # Build CSV file
    csv_stream = io.StringIO()
    writer = csv.writer(csv_stream)
    writer.writerow(["Action Agreed in Detail", "Owner", "Due Date"])
    for a in actions:
        writer.writerow([a.get("detail",""), a.get("owner",""), a.get("due","")])
    csv_b64 = base64.b64encode(csv_stream.getvalue().encode("utf-8")).decode("utf-8")

    return jsonify({
        "minutes_file": word_b64,
        "actions_file": csv_b64
    })
